from docx import Document
from docx.shared import Pt
from datetime import date
from pathlib import Path


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    if h.runs:
        h.runs[0].font.name = "Calibri"
    return h


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def main():
    doc = Document()

    add_heading(doc, "Informe de seguimiento de desarrollo - Proyecto Alcaldía", level=0)
    add_paragraph(doc, f"Fecha de elaboración: {date.today().strftime('%d/%m/%Y')}")

    add_heading(doc, "1. Seguimiento semanal", level=1)

    add_heading(doc, "Semana del 16 al 20 de marzo", level=2)
    add_paragraph(doc, "Se dio inicio al desarrollo del backend.")
    add_bullet(doc, "Estructura inicial de NestJS y configuración base del proyecto.")
    add_bullet(doc, "Configuración de TypeORM y conexión a PostgreSQL.")
    add_bullet(doc, "Verificación temprana de conectividad de base de datos.")

    add_heading(doc, "Semana del 23 al 27 de marzo", level=2)
    add_paragraph(doc, "Se implementó la capa inicial de autenticación y usuarios.")
    add_bullet(doc, "Endpoints de autenticación y registro.")
    add_bullet(doc, "DTOs y validaciones de entrada.")
    add_bullet(doc, "Guardias JWT para protección de rutas.")

    add_heading(doc, "Semana del 30 de marzo al 3 de abril", level=2)
    add_paragraph(doc, "Se consolidaron entidades y relaciones del dominio principal.")
    add_bullet(doc, "Modelado de proyectos y trámites.")
    add_bullet(doc, "Ajustes de relaciones TypeORM para evitar errores de propiedades.")
    add_bullet(doc, "Persistencia de vínculo trámite-proyecto (proyecto_id).")

    add_heading(doc, "Semana del 6 al 10 de abril", level=2)
    add_paragraph(doc, "Se habilitó la gestión documental en backend y frontend.")
    add_bullet(doc, "Subida, listado y descarga de documentos.")
    add_bullet(doc, "Almacenamiento de archivos en carpeta uploads.")
    add_bullet(doc, "Integración con formularios de creación de trámite/proyecto.")

    add_heading(doc, "Semana del 13 al 17 de abril", level=2)
    add_paragraph(doc, "Se agregó la generación de recibo bancario en PDF.")
    add_bullet(doc, "Endpoint para generar y descargar recibo PDF por trámite.")
    add_bullet(doc, "Integración de botón de descarga en detalle de trámite.")

    add_heading(doc, "Semana del 20 al 24 de abril", level=2)
    add_paragraph(doc, "Se realizaron ajustes funcionales de frontend y payloads.")
    add_bullet(doc, "Corrección de campos enviados desde registro de usuarios.")
    add_bullet(doc, "Alineación de nombres de campos frontend-backend.")
    add_bullet(doc, "Mejoras de navegación entre crear trámite y detalle de proyecto.")

    add_heading(doc, "Semana del 27 de abril al 1 de mayo", level=2)
    add_paragraph(doc, "Se fortaleció el módulo administrativo de locaciones e imágenes.")
    add_bullet(doc, "Correcciones de modal de imágenes y referencias DOM.")
    add_bullet(doc, "Carga de imágenes para locaciones desde panel admin.")
    add_bullet(doc, "Operaciones de listado, borrado y orden en galería.")

    add_heading(doc, "Semana del 4 al 8 de mayo", level=2)
    add_paragraph(doc, "Se incorporaron mejoras de portafolio proveedor y despliegue.")
    add_bullet(doc, "Cambio de enfoque de 'Agregar Proyecto' a 'Agregar servicios'.")
    add_bullet(doc, "Soporte de galería del proveedor y selección de servicios en proyectos.")
    add_bullet(doc, "Ajuste del script start:prod para compilar y luego ejecutar.")

    add_heading(doc, "2. Capturas recomendadas por apartado", level=1)

    add_heading(doc, "2.1 Portada / Autenticación", level=2)
    add_bullet(doc, "Pantalla de inicio de sesión centrada (public/iniciar-sesion/index.html).")
    add_bullet(doc, "Versión responsive en resolución móvil.")

    add_heading(doc, "2.2 Registro de usuario", level=2)
    add_bullet(doc, "Formulario completo de registro (public/registro/index.html).")
    add_bullet(doc, "Evidencia de validación o registro exitoso.")

    add_heading(doc, "2.3 Trámites y documentos", level=2)
    add_bullet(doc, "Crear trámite con selección de locación existente (public/crear-tramite/index.html).")
    add_bullet(doc, "Subida de documento en trámite/proyecto.")
    add_bullet(doc, "Detalle de trámite con botón de descarga de PDF (public/detalle-tramite/index.html).")

    add_heading(doc, "2.4 Proyectos", level=2)
    add_bullet(doc, "Detalle de proyecto mostrando trámites asociados (public/detalle-proyecto/index.html).")
    add_bullet(doc, "Crear proyecto con selección de servicio (public/crear-proyecto/index.html).")

    add_heading(doc, "2.5 Administración", level=2)
    add_bullet(doc, "Modal de imágenes en administración de locaciones (public/admin/admin-4.html).")
    add_bullet(doc, "Proceso de subida de imagen desde el modal.")

    add_heading(doc, "2.6 Proveedor / Portafolio", level=2)
    add_bullet(doc, "Listado de servicios y galería del proveedor (public/proveedor/proveedor-2.html).")
    add_bullet(doc, "Formulario para agregar servicio y evidencia de carga de imagen en galería.")

    add_heading(doc, "3. Recomendaciones de presentación", level=1)
    add_bullet(doc, "Nombrar capturas de forma secuencial: 01-login, 02-registro, 03-crear-tramite, etc.")
    add_bullet(doc, "Usar resolución horizontal (1280x720) para páginas completas.")
    add_bullet(doc, "Para modales, tomar captura con zoom suficiente para leer campos y botones.")
    add_bullet(doc, "Incluir una captura por flujo con estado final exitoso (guardado/subida/descarga).")

    output_dir = Path(__file__).resolve().parents[1] / "reportes"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / "informe_seguimiento_proyecto_alcaldia.docx"
    doc.save(output_file)

    print(f"Documento generado: {output_file}")


if __name__ == "__main__":
    main()
